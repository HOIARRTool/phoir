# streamlit_app.py
# -*- coding: utf-8 -*-

import os
import re
import json
import html
from io import BytesIO
from datetime import datetime, date, time
from typing import Any, Dict, List, Optional, Tuple
from textwrap import dedent

import pandas as pd
import requests
import streamlit as st
import gspread

from docx import Document
from docx.shared import Inches
from google.oauth2.service_account import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload


# =========================
# CONFIG / CONSTANTS
# =========================

LOGO_HA_URL = "https://raw.githubusercontent.com/HOIARRTool/appqtbi/main/messageImage_1763018963411.jpg"

UNIT_OPTIONS = [
    "PCU1",
    "PCU2",
    "PCU3",
    "PCU4",
    "PCU5",
]

# เก็บคอลัมน์เดิม + เพิ่มคอลัมน์ใหม่แบบ backward-compatible
SHEET_COLUMNS = [
    "record_id",
    "unit_name",                  # ใช้เก็บ “หน่วย” ที่ผู้ใช้เลือก
    "app_title",
    "event_date",                 # YYYY-MM-DD
    "event_time",                 # HH:MM

    # --- โครงสร้างเดิม (คงไว้เพื่อ compatibility) ---
    "process_step",               # เดิม: กระบวนการ / ปัจจุบันใช้เก็บ event_code
    "drug_name",                  # เดิม: ชื่อยา (ฟอร์มใหม่ไม่ใช้แล้ว)
    "severity_level",             # A-I หรือ 1-5
    "incident_detail",
    "timeline_text",
    "initial_correction",
    "rca_text",
    "rca_image_filename",
    "rca_image_drive_url",
    "development_plan",
    "created_at",
    "created_by",

    # --- คอลัมน์ใหม่ ---
    "incident_group",             # 4 กลุ่มใหญ่
    "event_code",                 # CPM201 / CPP101 / ...
    "event_topic",                # คำอธิบายหัวข้อ
    "severity_scheme",            # "A-I" / "1-5"
    "event_display",              # เก็บข้อความแสดงผลรวม code + topic
]

# รายการเหตุการณ์ตามกลุ่ม
INCIDENT_GROUP_OPTIONS = [
    "ความคลาดเคลื่อนทางยาและอาการไม่พึงประสงค์",
    "Patient Safety",
    "Personal Safety",
    "People Safety",
]

EVENT_CODE_MAP: Dict[str, List[Tuple[str, str]]] = {
    "ความคลาดเคลื่อนทางยาและอาการไม่พึงประสงค์": [
        ("CPM201", "Medication error : Prescribing (เกิดข้อผิดพลาด/อุบัติการณ์ในขั้นตอนการสั่งใช้ยา)"),
        ("CPM202", "Medication error : Transcribing (เกิดข้อผิดพลาด/อุบัติการณ์ในขั้นตอนการคัดลอกยา)"),
        ("CPM203", "Medication error : Pre-dispensing (เกิดข้อผิดพลาด/อุบัติการณ์ในขั้นตอนการจัดเตรียมจ่ายยา)"),
        ("CPM204", "Medication error : Dispensing (เกิดข้อผิดพลาด/อุบัติการณ์ในขั้นตอนการจ่ายยา)"),
        ("CPM205", "Medication error : Administration (เกิดข้อผิดพลาด/อุบัติการณ์ในขั้นตอนการให้ยา)"),
        ("CPM101", "แพ้ยาซ้ำ"),
        ("CPM102", "ไม่มี/ไม่ปฏิบัติตาม Guideline ของการใช้ High Alert Drug"),
        ("CPM103", "ผู้ป่วยมีภาวะแทรกซ้อนที่ป้องกันได้จากการได้รับยาความเสี่ยงสูง"),
        ("CPM104", "Mis selection of a strong potassium containing solution***"),
        ("CPM105", "แพ้ยา (ยกเว้น แพ้ยาซ้ำ)/ADE: Adverse Drug Events ที่มีความรุนแรงระดับ E ขึ้นไป"),
        ("CPM106", "ไม่มี/ไม่ปฏิบัติตาม Guideline ของการใช้ Fatal Drug"),
        ("CPM107", "ผู้ป่วยได้รับยาที่มีคู่ยาปฏิกิริยารุนแรง"),
        ("CPM206", "ไม่มี/ไม่ปฏิบัติตาม Guideline เกี่ยวกับ Look-Alike Sound-Alike Medication Names"),
        ("CPM207", "ผู้ป่วยได้รับยา ในกลุ่ม Look-Alike Sound-Alike Medication Names"),
        ("CPM208", "ไม่มี/ไม่ปฏิบัติตามมาตรฐาน หรือ Guideline ของการใช้ยา ยกเว้น HAD, Fatal drug, Look-Alike Sound-Alike, Antibiotics"),
        ("CPM301", "ไม่มี/ไม่ปฏิบัติตาม Guideline เกี่ยวกับ Medication Reconciliation"),
        ("CPM302", "ผู้ป่วยไม่ได้รับยาเดิมต่อเนื่องจากไม่ได้ทำ Medication Reconciliation"),
        ("CPM303", "ผู้ป่วยได้รับยาซ้ำซ้อนจากไม่ได้ทำ Medication Reconciliation"),
        ("CPM304", "ผู้ป่วยได้รับยาที่มีปฏิกิริยากันโดยไม่ได้ทำ Medication Reconciliation"),
        ("CPM401", "ไม่มี/ไม่ปฏิบัติตาม Guideline เกี่ยวกับ Rational Drug Use"),
        ("CPM402", "การใช้ยาปฏิชีวนะในโรคติดเชื้อที่ระบบการหายใจช่วงบนและหลอดลมอักเสบเฉียบพลันในผู้ป่วยนอก"),
        ("CPM403", "การใช้ยาปฏิชีวนะในโรคอุจจาระร่วงเฉียบพลัน"),
        ("CPM404", "การใช้ยาอย่างไม่สมเหตุผล (ยกเว้นยาปฏิชีวนะ)"),
        ("GOI207", "เกิดปัญหาด้านเวชภัณฑ์ยา เช่น ไม่มีแผนบริหารจัดการ/ไม่มีคุณภาพ/ไม่เพียงพอ/หมดอายุหรืออยู่ในสภาพไม่พร้อมใช้งาน"),
    ],
    "Patient Safety": [
        ("CPL202", "สิ่งส่งตรวจไม่ถูกต้อง ไม่เหมาะสม หรือไม่มีสิ่งส่งตรวจ"),
        ("CPP101", "Patient Identification"),
        ("CPP201", "การรายงานอาการ หรือสื่อสารข้อมูลเกี่ยวกับผู้ป่วยไม่เหมาะสม/ไม่ครบถ้วน"),
        ("CPP204", "การสื่อสารหรือส่งต่อข้อมูลการรักษาพยาบาลผู้ป่วยผิดพลาด เช่น ไม่สื่อสาร/สื่อสารผิด/สื่อสารไม่ครบถ้วน/สื่อสารล่าช้า"),
        ("CPP206", "เกิดความผิดพลาดในการรักษาพยาบาลซึ่งมีสาเหตุมาจาก Verbal or Telephone Order/Communication"),
        ("CPP303", "(Patient Assessment) ผู้ป่วยไม่ได้รับการประเมิน/ประเมินผิด/ประเมินไม่ครบถ้วน ตามเกณฑ์ อาการหรือการดำเนินโรค"),
        ("CPP405", "ตกเตียง/fall"),
        ("CPP601", "ผู้ป่วยที่จำเป็นต้องส่งต่อเพื่อการรักษา ไม่ได้รับการส่งต่อหรือส่งต่อได้ในช่วงเวลาไม่เหมาะสม"),
        ("CPP602", "มีภาวะแทรกซ้อนหรือเหตุการณ์ไม่พึงประสงค์ที่ป้องกันได้ระหว่างส่งต่อ"),
    ],
    "Personal Safety": [
        ("GPE101", "อันตรายจากโครงสร้างอาคารสถานที่และสิ่งแวดล้อมเชิงกายภาพ เช่น แสง เสียง ฝุ่นละออง มีเชื้อรา เป็นต้น"),
        ("GPE205", "เกิดปัญหาด้านการจัดการสภาพแวดล้อมในการให้บริการ เช่น ไม่มีป้ายให้คำแนะนำ/บอกทาง, ไม่มีทางหนีไฟหรือมีแต่ไม่พร้อมใช้/มีสิ่งกีดขวาง, ลิฟต์ขัดข้อง มีคนติดในลิฟต์ หรือ ลิฟต์ไม่พร้อมใช้งาน/ชำรุด/ติดค้าง"),
        ("GPE206", "เกิดปัญหาด้านการควบคุมสิ่งแวดล้อมในสถานที่ทำงาน เช่น ระบบน้ำอุปโภค-บริโภคไม่เพียงพอ/ไม่พร้อมใช้, ระบบไฟฟ้าไม่เพียงพอ ไม่พร้อมใช้/ดับ/ช็อต/กระพริบ, การบำบัดน้ำเสีย/กำจัดขยะ ไม่ถูกวิธี/ไม่ได้มาตรฐาน"),
        ("GPE303", "บุคลากรได้รับภัยคุกคามหรือถูกทำร้ายทางวาจาจากผู้ป่วยและญาติหรือบุคคลภายนอก"),
        ("GPE304", "บุคลากรได้รับภัยคุกคามหรือถูกทำร้ายทางกายจากผู้ป่วยและญาติหรือบุคคลภายนอก"),
        ("GPE305", "เกิดกรณีความไม่สงบในสถานพยาบาล เช่น เมาสุราอาละวาด"),
        ("GPI101", "บุคลากรถูกวัสดุอุปกรณ์มีคมทิ่มตำ"),
        ("GPI102", "บุคลากรสัมผัสเลือดหรือสารคัดหลั่งบริเวณเยื่อบุหรือผิวหนังที่มีแผล (mucous membrane and non-intact skin exposure to blood and body fluid)"),
        ("GPI201", "บุคลากรติดเชื้อที่แพร่ทางอากาศ (airborne transmission) จากการปฏิบัติงาน ได้แก่ วัณโรค หัด และอีสุกอีใส"),
        ("GPL101", "อุปกรณ์บนรถพยาบาลไม่พร้อมใช้ ไม่เหมาะสมและไม่ปลอดภัยสำหรับการส่งต่อผู้ป่วย"),
        ("GPL104", "เกิดอุบัติเหตุของรถพยาบาลระหว่างปฏิบัติหน้าที่"),
        ("GPL106", "พนักงานขับรถไม่ปฏิบัติตามแนวทางความปลอดภัยของรถบริการการแพทย์ฉุกเฉิน และรถพยาบาล เช่น ขับรถเร็วเกินกว่ากำหนด"),
        ("GPM104", "เจ้าหน้าที่มีภาวะเครียดจากการทำงาน"),
        ("GPM203", "เกิดเรื่องร้องเรียนจากการบริการทางการแพทย์"),
        ("GPM204", "เกิดเรื่องร้องเรียนทั่วไป ซึ่งไม่เกี่ยวกับการบริการทางการแพทย์"),
        ("GPP203", "บุคลากรเกิดโรคจากการทำงาน ซึ่งมีสาเหตุจาก Physical Hazard"),
        ("GPP212", "บุคลากรเกิดโรคจากการทำงานเกี่ยวกับโครงร่างกระดูกและกล้ามเนื้อ ซึ่งมีสาเหตุจาก Biomechanical Hazard"),
        ("GPS105", "เกิดอุบัติการณ์การละเมิดความเป็นส่วนตัว (Privacy) ของข้อมูลส่วนบุคคลของบุคลากรหรือนักศึกษาของสถานพยาบาล ที่ไม่ใช่อุบัติการณ์ด้านความมั่นคงปลอดภัยไซเบอร์"),
        ("GPS106", "เกิดอุบัติการณ์ความละเมิดความเป็นส่วนตัว (Privacy) ของข้อมูลส่วนบุคคลของผู้ป่วย/ผู้รับบริการ หรือบุคคลภายนอก ที่ไม่ใช่อุบัติการณ์ด้านความมั่นคงปลอดภัยไซเบอร์"),
        ("GPS203", "บุคลากรใช้สื่อสังคมออนไลน์ไม่เหมาะสม เกิดผลกระทบทางลบต่อตนเอง บุคลากรคนอื่น สถานพยาบาล ผู้ป่วย/ผู้รับบริการ หรือบุคคลภายนอก"),
    ],
    "People Safety": [
        ("GOI101", "เกิดปัญหาด้าน Hardware/อุปกรณ์คอมพิวเตอร์ เช่น ไม่มีแผนบริหารจัดการ/ไม่เพียงพอ/ไม่พร้อมใช้/ใช้ไม่ตรงวัตถุประสงค์/ใช้ผิดวิธี-เทคนิค"),
        ("GOI102", "เกิดปัญหาด้าน Network & Security เช่น ไม่พร้อมใช้/ระบบล่ม/มีการเข้าถึงโดยผู้ไม่มีสิทธิ์"),
        ("GOI107", "โปรแกรม/ระบบสารสนเทศทางการแพทย์ เช่น HIS, LIS, โปรแกรมระบบยา ล่ม/ไม่สามารถใช้งานได้ นานเกินกว่าระยะเวลาที่ประกันเวลาไว้"),
        ("GOI204", "เกิดปัญหาด้านอุปกรณ์เทคโนโลยีทางการแพทย์/เครื่องมือ-อุปกรณ์การแพทย์ที่ไม่ใช่เครื่องมือ-อุปกรณ์ผ่าตัด (Error of Medical device) เช่น ไม่มีแผนบริหารจัดการ/ไม่เพียงพอ/ไม่พร้อมใช้/ใช้ไม่ตรงวัตถุประสงค์/ใช้ผิดวิธี-เทคนิค"),
        ("GOI206", "เกิดปัญหาด้านเวชภัณฑ์ทางการแพทย์/เวชภัณฑ์ที่ไม่ใช่ยา เช่น ไม่มีแผนบริหารจัดการ/ไม่มีคุณภาพ/ไม่เพียงพอ/หมดอายุหรืออยู่ในสภาพไม่พร้อมใช้งาน"),
        ("GOS201", "อาคารสถานที่/พื้นที่ให้บริการ ไม่เหมาะสม/ไม่ปลอดภัย/ไม่ถูกสุขลักษณะ"),
        ("GOS202", "ห้องน้ำหรือห้องสุขาไม่พร้อมใช้ (เช่น ชำรุด/กดชักโครกไม่ลง/ส้วมเต็ม/ไม่พอใช้) หรือไม่สะดวกต่อผู้พิการ"),
    ],
}

SEVERITY_OPTIONS_AI = list("ABCDEFGHI")
SEVERITY_OPTIONS_PEOPLE = ["1", "2", "3", "4", "5"]

SEVERITY_DESC_AI = {
    "A": "(เกิดที่นี่) มีโอกาสเกิดเหตุการณ์และค้นพบได้ด้วยตัวเอง สามารถปรับแก้ไขได้ ไม่ส่งผลกระทบถึงผู้อื่นและผู้ป่วยหรือบุคลากร ไม่เกิดความรุนแรง (No Harm)",
    "B": "(เกิดที่ไกล) เกิดเหตุการณ์/ความผิดพลาดขึ้นแล้วโดยส่งต่อเหตุการณ์/ความผิดพลาดนั้นไปที่ผู้อื่นแต่สามารถตรวจพบและแก้ไขได้ โดยยังไม่มีผลกระทบใดๆ ถึงผู้ป่วยหรือบุคลากร ไม่เกิดความรุนแรง (No Harm)",
    "C": "(เกิดกับใคร) เกิดเหตุการณ์/ความผิดพลาดขึ้นและมีผลกระทบถึงผู้ป่วยหรือบุคลากร แต่ไม่เกิดอันตรายหรือเสียหาย เกิดความรุนแรงน้อย (Low Harm)",
    "D": "(ให้ระวัง) เกิดความผิดพลาดขึ้น มีผลกระทบถึงผู้ป่วยหรือบุคลากร ต้องให้การดูแลเฝ้าระวังเป็นพิเศษว่าจะไม่เป็นอันตราย เกิดความรุนแรงน้อย (Low Harm)",
    "E": "(ต้องรักษา) เกิดความผิดพลาดขึ้น มีผลกระทบถึงผู้ป่วยหรือบุคลากร เกิดอันตรายชั่วคราวที่ต้องแก้ไข/รักษาเพิ่มมากขึ้น เกิดความรุนแรงปานกลาง (Moderate Harm)",
    "F": "(เยียวยานาน) เกิดความผิดพลาดขึ้น มีผลกระทบที่ต้องใช้เวลาแก้ไขนานกว่าปกติหรือเกินกำหนด ผู้ป่วยหรือบุคลากร ต้องรักษา/นอนโรงพยาบาลนานขึ้น เกิดความรุนแรงปานกลาง (Moderate Harm)",
    "G": "(ต้องพิการ) เกิดความผิดพลาดถึงผู้ป่วยหรือบุคลากร ทำให้เกิดความพิการถาวร หรือมีผลกระทบทำให้เสียชื่อเสียง/ความเชื่อถือและ/หรือมีการร้องเรียน เกิดความรุนแรงมาก (Severe Harm)",
    "H": "(ต้องการปั๊ม) เกิดความผิดพลาด ถึงผู้ป่วยหรือบุคลากร มีผลทำให้ต้องทำการช่วยชีวิต ต้องดูแลต่อเนื่องตลอดชีวิต ช่วยเหลือตัวเองไม่ได้ หรือกรณีทำให้เสียชื่อเสียงและ/หรือมีการเรียกร้องค่าเสียหายจากโรงพยาบาล เกิดความรุนแรงมาก (Severe Harm)",
    "I": "(จำใจลา) เกิดความผิดพลาด ถึงผู้ป่วยหรือบุคลากร เป็นสาเหตุทำให้เสียชีวิต เสียชื่อเสียงโดยมีการฟ้องร้องทางศาล/สื่อ เสียชีวิต (Death)",
}

SEVERITY_DESC_PEOPLE = {
    "1": "เกิดความผิดพลาดขึ้นแต่ไม่มีผลกระทบต่อผลสำเร็จหรือวัตถุประสงค์ของการดำเนินงาน (ดำเนินงานสำเร็จตามแผนได้มากกว่า 90%) หรือ ทำให้เกิดความล่าช้าของโครงการ ไม่เกิน 1.5 เดือน หรือ ผลกระทบด้านการเงินมีมูลค่าน้อยกว่า 10,000 บาท",
    "2": "เกิดความผิดพลาดขึ้นแล้ว โดยมีผลกระทบ (ที่ควบคุมได้) ต่อผลสำเร็จหรือวัตถุประสงค์ของการดำเนินงาน (ดำเนินงานสำเร็จตามแผนได้ 81 - 90%) หรือ ทำให้เกิดความล่าช้าของโครงการ มากกว่า 1.5 - 3 เดือน หรือ ผลกระทบด้านการเงินมีมูลค่า 10,001 – 100,000 บาท",
    "3": "เกิดความผิดพลาดขึ้นแล้ว และมีผลกระทบ (ที่ต้องทำการแก้ไข) ต่อผลสำเร็จหรือวัตถุประสงค์ของการดำเนินงาน (ดำเนินงานสำเร็จตามแผนได้ 71 - 80%) หรือ ทำให้เกิดความล่าช้าของโครงการ มากกว่า 3 – 4.5 เดือน หรือ ผลกระทบด้านการเงินมีมูลค่า 100,001 – 500,000 บาท",
    "4": "เกิดความผิดพลาดขึ้นแล้ว และทำให้การดำเนินงานไม่บรรลุผลสำเร็จตามเป้าหมาย (ดำเนินงานสำเร็จตามแผนได้ 60 - 70%) หรือ ทำให้เกิดความล่าช้าของโครงการ มากกว่า 4.5 - 6 เดือน หรือ ผลกระทบด้านการเงินมีมูลค่า 500,001 – 10,000,000 บาท",
    "5": "เกิดความผิดพลาดขึ้นแล้ว และมีผลให้การดำเนินงานไม่บรรลุผลสำเร็จตามเป้าหมาย (ดำเนินงานสำเร็จตามแผนได้น้อยกว่า 60%) ทำให้ภารกิจขององค์กรเสียหายอย่างร้ายแรง หรือ ทำให้เกิดความล่าช้าของโครงการ มากกว่า 6 เดือน หรือ ผลกระทบด้านการเงินมีมูลค่ามากกว่า 10,000,000 บาท",
}


# =========================
# PAGE SETUP
# =========================

st.set_page_config(
    page_title="PHOIR",
    page_icon="🏡",
    layout="wide",
)


# =========================
# HELPER: READ CONFIG (ENV ONLY for Render)
# =========================

def _get_env(
    key: str,
    default: Optional[str] = None,
    aliases: Optional[List[str]] = None,
) -> Optional[str]:
    keys = [key] + (aliases or [])
    for k in keys:
        v = os.getenv(k)
        if v is not None and str(v).strip() != "":
            return str(v).strip()
    return default


def get_app_config() -> Dict[str, Any]:
    app_title = _get_env("APP_TITLE", "PHOIR_DEMO")
    unit_name = _get_env("UNIT_NAME", "unknown-unit")
    login_user = _get_env("APP_LOGIN_USERNAME", "")
    login_pass = _get_env("APP_LOGIN_PASSWORD", "")

    gsheet_url = _get_env("GSHEET_URL", "")
    worksheet_name = _get_env("GSHEET_WORKSHEET", "PHOIR_DEMO", aliases=["GHEET_WORKSHEET"])

    gcp_sa_json = _get_env("GCP_SERVICE_ACCOUNT_JSON", "", aliases=["GSHEET_CREDENTIALS_JSON"])
    gemini_api_key = _get_env("GEMINI_API_KEY", "")
    gdrive_folder_id = _get_env("GDRIVE_FOLDER_ID", "")

    return {
        "APP_TITLE": app_title,
        "UNIT_NAME": unit_name,
        "APP_LOGIN_USERNAME": login_user,
        "APP_LOGIN_PASSWORD": login_pass,
        "GSHEET_URL": gsheet_url,
        "GSHEET_WORKSHEET": worksheet_name,
        "GCP_SERVICE_ACCOUNT_JSON": gcp_sa_json,
        "GEMINI_API_KEY": gemini_api_key,
        "GDRIVE_FOLDER_ID": gdrive_folder_id,
    }


CFG = get_app_config()


# =========================
# STYLING
# =========================

st.markdown(
    """
<style>
.block-container { padding-top: 1rem; }
.small-muted { color: #6b7280; font-size: 1rem; }
.card {
    border: 1px solid #e5e7eb;
    border-radius: 14px;
    padding: 14px;
    background: #ffffff;
}
.section-title {
    font-size: 1.05rem;
    font-weight: 700;
    margin-bottom: .5rem;
}

/* ===== Login shell / hero ===== */
.login-shell-wrapper hr { display: none !important; }


.login-shell::before {
    display: none !important;
    content: none !important;
}

.login-hero {
    background: rgba(255,255,255,0.90);
    border: 1px solid rgba(148,163,184,0.20);
    border-radius: 18px;
    padding: 14px 14px 10px 14px;
    margin-bottom: 10px;
    box-shadow: none !important;
    backdrop-filter: none !important;
}

.logo-row {
    display: flex;
    justify-content: center;
    align-items: flex-end;
    gap: 18px;
    margin-bottom: 8px;
}
.logo-left-wrap,
.logo-right-wrap {
    display: flex;
    align-items: center;
    justify-content: center;
}
.logo-left-wrap {
    width: 108px;   
    height: 86px;
}
.logo-right-wrap {
    width: 96px;
    height: 86px;
}
.logo-left-wrap img {
    max-width: 88px;
    max-height: 72px;
    object-fit: contain;
}
.logo-right-wrap img {
    max-width: 74px;
    max-height: 72px;
    object-fit: contain;
}

.login-title-center {
    text-align: center;
}
.login-title-center h1 {
    margin: 0;
    font-size: 1.95rem;
    line-height: 1.1;
    color: #2b2d42;
    font-weight: 800;
    letter-spacing: 0.2px;
}
.login-title-center .subtitle {
    margin-top: 6px;
    color: #4b5563;
    font-size: 1.02rem;
    font-weight: 500;
}

.login-card-box {
    background: rgba(255,255,255,0.92);
    border: 1px solid rgba(148,163,184,0.22);
    border-radius: 18px;
    padding: 14px;
    box-shadow: none !important;
    backdrop-filter: none !important;
}
.login-card-box h3 {
    margin-top: 0.1rem !important;
    margin-bottom: 0.35rem !important;
    font-size: 1.5rem !important;
}
.login-card-box .login-head-note {
    color: #6b7280;
    font-size: 1rem;
    margin-bottom: 0.5rem;
}
.login-card-box [data-testid="stTextInput"] label,
.login-card-box [data-testid="stTextInputRootElement"] label {
    font-weight: 600;
}
.login-card-box [data-testid="stTextInput"] > div > div input {
    border-radius: 10px;
}
.login-card-box [data-testid="stButton"] button {
    border-radius: 10px;
    font-weight: 600;
}

.login-info-box {
    background: rgba(255,255,255,0.92);
    border: 1px solid rgba(148,163,184,0.22);
    border-radius: 18px;
    padding: 12px 14px;
    box-shadow: none !important;
    backdrop-filter: none !important;
}

.login-badge {
    display: inline-block;
    padding: 4px 10px;
    border-radius: 999px;
    background: #eef2ff;
    border: 1px solid #c7d2fe;
    color: #4338ca;
    font-weight: 700;
    font-size: 1rem;
    margin-bottom: 0.55rem;
}
.login-paragraph {
    margin-bottom: 0.55rem;
    color: #1f2937;
    line-height: 1.55;
    font-size: 1.2rem;
}
.feature-list {
    display: grid;
    gap: 8px;
    margin-bottom: 0.6rem;
}
.feature-item {
    border: 1px solid #dbeafe;
    background: #f8fbff;
    border-radius: 12px;
    padding: 8px 10px;
}
.feature-title {
    font-weight: 700;
    color: #1e3a8a;
    margin-bottom: 2px;
    font-size: 1.2rem;
}
.feature-desc {
    color: #475569;
    line-height: 1.45;
    font-size: 1.2rem;
}
.quote-box {
    border: 1px dashed #93c5fd;
    background: #f8fbff;
    border-radius: 12px;
    padding: 9px 10px;
    color: #1e293b;
    line-height: 1.5;
    font-size: 1.2rem;
}

/* Compact spacing on login widgets */
.login-card-box .stTextInput, .login-card-box .stButton {
    margin-bottom: 0.15rem;
}
.login-card-box div[data-testid="stCaptionContainer"] {
    margin-top: -2px;
}

/* Entry form helpers */
.helper-box {
    border: 1px solid #e5e7eb;
    border-radius: 12px;
    padding: 10px 12px;
    background: #fafafa;
}
.helper-title {
    font-weight: 700;
    margin-bottom: 6px;
    color: #111827;
}
.tiny {
    font-size: 1rem;
    color: #6b7280;
}

@media (max-width: 1100px) {
    .login-title-center h1 { font-size: 1.72rem; }
    .login-title-center .subtitle { font-size: 1.2rem; }
    .logo-left-wrap { width: 94px; height: 78px; }
    .logo-right-wrap { width: 84px; height: 78px; }
    .logo-left-wrap img { max-width: 76px; max-height: 64px; }
    .logo-right-wrap img { max-width: 64px; max-height: 64px; }
}

@media (max-width: 768px) {
    .login-shell { padding: 12px; border-radius: 16px; }
    .login-hero { padding: 12px 10px 8px 10px; }
    .logo-row { gap: 10px; margin-bottom: 6px; }
    .logo-left-wrap, .logo-right-wrap { height: 64px; }
    .logo-left-wrap { width: 78px; }
    .logo-right-wrap { width: 70px; }
    .logo-left-wrap img { max-width: 62px; max-height: 56px; }
    .logo-right-wrap img { max-width: 52px; max-height: 56px; }
    .login-title-center h1 { font-size: 1.36rem; }
    .login-title-center .subtitle { font-size: 1rem; }
    .login-info-box, .login-card-box { border-radius: 14px; }
    .login-paragraph { font-size: 1.2rem; }
}
</style>
    """,
    unsafe_allow_html=True,
)


# =========================
# LOGIN
# =========================

def ensure_auth_state():
    if "authenticated" not in st.session_state:
        st.session_state.authenticated = False
    if "login_username" not in st.session_state:
        st.session_state.login_username = ""
    if "show_fishbone_preview" not in st.session_state:
        st.session_state.show_fishbone_preview = False


def render_login_header_hero():
    html_block = (
        "<div class='login-hero'>"
        "<div class='logo-row'>"
        f"<div class='logo-left-wrap'><img src='{LOGO_HA_URL}' alt='HA Logo'></div>"
        "</div>"
        "<div class='login-title-center'>"
        f"<h1> {html.escape(CFG['APP_TITLE'])}</h1>"
        "<div class='title'><strong>P</strong>rimary <strong>H</strong>ealthcare <strong>O</strong>ccurence/<strong>I</strong>ncident <strong>R</strong>eport: บันทึกอุบัติการณ์ในสถานพยาบาลปฐมภูมิ</div>"
        "</div>"
        "</div>"
    )
    st.markdown(html_block, unsafe_allow_html=True)


def render_login_info_panel():
    html_block = (
        "<div class='login-info-box'>"
        "<div class='login-badge'>อ้างอิงมาตรฐานสถานพยาบาลปฐมภูมิ (สรพ.) ฉบับที่ 1</div>"

        "<div class='login-paragraph'>"
        "ระบบนี้ถูกออกแบบมาเพื่อยกระดับความปลอดภัยในหน่วยบริการปฐมภูมิ <br>"
        "สนับสนุนการรายงานอุบัติการณ์และรวบรวมเป็นข้อมูลเชิงลึก <br>"
        "ให้คนหน้างานริหารความเสี่ยงและพัฒนาคุณภาพได้ด้วยตนเอง <br>"
        "โดยต่อยอดจากบริบทและทรัพยากรที่มีอยู่จริงในพื้นที่"
        "</div>"

        "<div class='login-paragraph'>"
        "ขับเคลื่อนคุณภาพตามมาตรฐานสถานพยาบาลปฐมภูมิ (สรพ.) ได้แก่<br>"
        "<strong>ตอนที่ 1.6 ข. การบริหารความเสี่ยงและจัดการความปลอดภัย</strong> และ <br>"
        "<strong>ตอนที่ 4.3 การจัดการด้านยาและวัคซีน</strong> โดยมีเครื่องมือรองรับ<br>"
        "การจัดการความคลาดเคลื่อนทางยา (Medication Error)<br> "
        "และเหตุการณ์ไม่พึงประสงค์ อย่างเป็นระบบ ตามมาตรฐาน<br>"
        "ข้อ 4.3 (ข.4) มีการเก็บรวบรวมข้อมูลความคลาดเคลื่อนทางยาและเหตุการณ์ไม่พึงประสงค์"
        "</div>"
        "<div class='login-info-box'>"
        "<div class='login-badge'>เปลี่ยนข้อมูล สู่ความปลอดภัยที่ยั่งยืน</div>"

        "<div class='feature-list'>"
        "<div class='feature-item'>"
        "<div class='feature-title'>📝 บันทึกง่าย ครอบคลุม (Record with Ease)</div>"
        "<div class='feature-desc'>รายงานความเสี่ยงได้รวดเร็ว ไม่เพิ่มภาระงาน</div>"
        "</div>"

        "<div class='feature-item'>"
        "<div class='feature-title'>📊 วิเคราะห์ได้ด้วยตัวเอง (Local Data Analysis)</div>"
        "<div class='feature-desc'>มี RCA assistant ช่วยวิเคราะห์รายอุบัติการณ์ได้ตั้งแต่ตอนรายงาน<br>"
        "พร้อมคืนข้อมูลสู่คนทำงาน เข้าถึง Dashboard ประเมินสถานการณ์ได้ทุกเวลา</div>"
        "</div>"

        "<div class='feature-item'>"
        "<div class='feature-title'>🔄 พัฒนาอย่างต่อเนื่อง (Continuous Improvement)</div>"
        "<div class='feature-desc'>ขับเคลื่อนการแก้ปัญหาด้วยข้อมูล (Data-Driven) ป้องกันการเกิดซ้ำอย่างตรงจุด</div>"
        "</div>"
        "</div>"

        "<div class='quote-box'>"
        "<strong>\"การรายงานไม่ได้มีไว้เพื่อจับผิด แต่มีไว้เพื่อเรียนรู้\"</strong><br>"
        "ทุกข้อมูลคือโอกาส มาร่วมสร้าง <strong>วัฒนธรรมความปลอดภัย (Safety Culture)</strong> ที่แข็งแกร่ง<br>"
        "เพื่อยกระดับมาตรฐานการดูแลพี่น้องประชาชนในชุมชนของคุณ"
        "</div>"
        "</div>"
    )
    st.markdown(html_block, unsafe_allow_html=True)


def render_login():
    ensure_auth_state()

    st.markdown('<div class="login-shell-wrapper">', unsafe_allow_html=True)
    st.markdown('<div class="login-shell">', unsafe_allow_html=True)

    render_login_header_hero()

    left, right = st.columns([0.9, 1.7], gap="large")

    with left:
        st.markdown('<div class="login-card-box">', unsafe_allow_html=True)
        st.markdown("### 🔐 เข้าสู่ระบบ")
        st.markdown(f"<div class='login-head-note'>หน่วยงานระบบ: <b>{html.escape(CFG['UNIT_NAME'])}</b></div>", unsafe_allow_html=True)

        username = st.text_input("ชื่อผู้ใช้", key="login_user_input", placeholder="กรอกชื่อผู้ใช้")
        password = st.text_input("รหัสผ่าน", type="password", key="login_pass_input", placeholder="กรอกรหัสผ่าน")

        if st.button("เข้าสู่ระบบ", use_container_width=True):
            expected_user = CFG["APP_LOGIN_USERNAME"]
            expected_pass = CFG["APP_LOGIN_PASSWORD"]

            # Dev bypass
            if not expected_user or not expected_pass:
                st.session_state.authenticated = True
                st.session_state.login_username = username or "dev-user"
                st.warning("ยังไม่ได้ตั้งค่า APP_LOGIN_USERNAME / APP_LOGIN_PASSWORD ใน ENV → เข้าแบบ dev mode")
                st.rerun()

            if username == expected_user and password == expected_pass:
                st.session_state.authenticated = True
                st.session_state.login_username = username
                st.success("เข้าสู่ระบบสำเร็จ ✅")
                st.rerun()
            else:
                st.error("ชื่อผู้ใช้หรือรหัสผ่านไม่ถูกต้อง")

        st.markdown("</div>", unsafe_allow_html=True)

    with right:
        render_login_info_panel()

    st.markdown("</div>", unsafe_allow_html=True)
    st.markdown("</div>", unsafe_allow_html=True)


# =========================
# GOOGLE API (Sheets + Drive)
# =========================

@st.cache_resource(show_spinner=False)
def get_google_credentials():
    sa_json_str = CFG["GCP_SERVICE_ACCOUNT_JSON"]
    if not sa_json_str:
        raise ValueError("ไม่พบ GCP_SERVICE_ACCOUNT_JSON ใน Environment Variables")

    try:
        creds_dict = json.loads(sa_json_str)
    except json.JSONDecodeError as e:
        raise ValueError(f"GCP_SERVICE_ACCOUNT_JSON ไม่ใช่ JSON ที่ถูกต้อง: {e}")

    scopes = [
        "https://www.googleapis.com/auth/spreadsheets",
        "https://www.googleapis.com/auth/drive",
    ]
    creds = Credentials.from_service_account_info(creds_dict, scopes=scopes)
    return creds


@st.cache_resource(show_spinner=False)
def get_gspread_client():
    creds = get_google_credentials()
    return gspread.authorize(creds)


@st.cache_resource(show_spinner=False)
def get_drive_service():
    creds = get_google_credentials()
    return build("drive", "v3", credentials=creds, cache_discovery=False)


# =========================
# GOOGLE SHEETS
# =========================

@st.cache_resource(show_spinner=False)
def get_worksheet():
    gsheet_url = CFG["GSHEET_URL"]
    worksheet_name = CFG["GSHEET_WORKSHEET"]

    if not gsheet_url:
        raise ValueError("ไม่พบ GSHEET_URL ใน Environment Variables")

    client = get_gspread_client()
    sh = client.open_by_url(gsheet_url)

    try:
        ws = sh.worksheet(worksheet_name)
    except gspread.WorksheetNotFound:
        ws = sh.add_worksheet(title=worksheet_name, rows=1000, cols=80)

    header = ws.row_values(1)
    if not header:
        ws.append_row(SHEET_COLUMNS, value_input_option="USER_ENTERED")
    else:
        missing_cols = [c for c in SHEET_COLUMNS if c not in header]
        if missing_cols:
            all_vals = ws.get_all_values()
            if all_vals:
                df_old = pd.DataFrame(all_vals[1:], columns=all_vals[0])
            else:
                df_old = pd.DataFrame(columns=[])

            for col in SHEET_COLUMNS:
                if col not in df_old.columns:
                    df_old[col] = ""

            df_old = df_old[SHEET_COLUMNS]

            ws.clear()
            ws.append_row(SHEET_COLUMNS, value_input_option="USER_ENTERED")
            if not df_old.empty:
                ws.append_rows(
                    df_old.fillna("").astype(str).values.tolist(),
                    value_input_option="USER_ENTERED",
                )

    return ws


def append_record_to_sheet(record: Dict[str, Any]) -> None:
    ws = get_worksheet()
    row = []
    for col in SHEET_COLUMNS:
        val = record.get(col, "")
        row.append("" if val is None else str(val))
    ws.append_row(row, value_input_option="USER_ENTERED")


@st.cache_data(show_spinner=False, ttl=30)
def load_sheet_df() -> pd.DataFrame:
    ws = get_worksheet()
    records = ws.get_all_records()
    if not records:
        return pd.DataFrame(columns=SHEET_COLUMNS)

    df = pd.DataFrame(records)
    for c in SHEET_COLUMNS:
        if c not in df.columns:
            df[c] = ""

    # normalize legacy rows
    if "incident_group" not in df.columns:
        df["incident_group"] = ""
    if "event_code" not in df.columns:
        df["event_code"] = ""
    if "event_topic" not in df.columns:
        df["event_topic"] = ""
    if "event_display" not in df.columns:
        df["event_display"] = ""
    if "severity_scheme" not in df.columns:
        df["severity_scheme"] = ""

    # เติม fallback จาก process_step ให้แถวเก่า
    df["event_code"] = df["event_code"].astype(str)
    df["event_topic"] = df["event_topic"].astype(str)
    df["event_display"] = df["event_display"].astype(str)
    df["process_step"] = df["process_step"].astype(str)

    mask_no_code = df["event_code"].str.strip().eq("")
    df.loc[mask_no_code, "event_code"] = df.loc[mask_no_code, "process_step"]

    mask_no_disp = df["event_display"].str.strip().eq("")
    df.loc[mask_no_disp, "event_display"] = df.loc[mask_no_disp, "process_step"]

    return df[SHEET_COLUMNS]


# =========================
# GOOGLE DRIVE UPLOAD (RCA IMAGE)
# =========================

def upload_rca_image_to_drive(uploaded_file: Any, record_id: str) -> Dict[str, str]:
    if uploaded_file is None:
        return {"file_id": "", "file_name": "", "file_url": ""}

    folder_id = str(CFG.get("GDRIVE_FOLDER_ID", "") or "").strip()
    if not folder_id:
        raise ValueError("ยังไม่ได้ตั้งค่า GDRIVE_FOLDER_ID ใน Environment Variables")

    drive = get_drive_service()

    original_name = getattr(uploaded_file, "name", "attachment.png")
    mime_type = getattr(uploaded_file, "type", None) or "application/octet-stream"
    safe_name = f"{record_id}_{original_name}"

    file_metadata = {"name": safe_name, "parents": [folder_id]}

    media = MediaIoBaseUpload(
        BytesIO(uploaded_file.getvalue()),
        mimetype=mime_type,
        resumable=False,
    )

    created = drive.files().create(
        body=file_metadata,
        media_body=media,
        fields="id,name",
        supportsAllDrives=True,
    ).execute()

    file_id = created.get("id", "")
    file_name = created.get("name", safe_name)
    file_url = f"https://drive.google.com/file/d/{file_id}/view" if file_id else ""

    return {"file_id": file_id, "file_name": file_name, "file_url": file_url}


# =========================
# HELPERS: EVENT CODES / SEVERITY
# =========================

def event_code_options_for_group(group_name: str) -> List[str]:
    items = EVENT_CODE_MAP.get(group_name, [])
    opts = [f"{code} | {topic}" for code, topic in items]
    opts.append("อื่น ๆ | ระบุรหัส/หัวข้อเอง")
    return opts


def parse_event_code_option(selected: str) -> Tuple[str, str]:
    s = str(selected or "").strip()
    if not s:
        return "", ""
    if s.startswith("อื่น ๆ"):
        return "OTHER", "ระบุรหัส/หัวข้อเอง"
    if "|" in s:
        code, topic = s.split("|", 1)
        return code.strip(), topic.strip()
    return s.strip(), ""


def current_severity_scheme(group_name: str) -> str:
    return "1-5" if group_name == "People Safety" else "A-I"


def severity_options_for_group(group_name: str) -> List[str]:
    return SEVERITY_OPTIONS_PEOPLE if current_severity_scheme(group_name) == "1-5" else SEVERITY_OPTIONS_AI


def severity_description(level: str, group_name: str) -> str:
    if current_severity_scheme(group_name) == "1-5":
        return SEVERITY_DESC_PEOPLE.get(str(level), "")
    return SEVERITY_DESC_AI.get(str(level), "")


def render_severity_guide(group_name: str):
    scheme = current_severity_scheme(group_name)
    with st.expander(f"📘 คำอธิบายระดับความรุนแรง ({scheme})", expanded=False):
        if scheme == "1-5":
            data = [{"ระดับ": k, "รายละเอียด": v} for k, v in SEVERITY_DESC_PEOPLE.items()]
        else:
            data = [{"ระดับ": k, "รายละเอียด": v} for k, v in SEVERITY_DESC_AI.items()]
        st.dataframe(pd.DataFrame(data), use_container_width=True, hide_index=True)


# =========================
# DOCX EXPORT (BEFORE SAVE)
# =========================

def build_docx_report_bytes(uploaded_rca_image: Optional[Any] = None) -> bytes:
    doc = Document()

    doc.add_heading("รายงานอุบัติการณ์ / RCA (ก่อนบันทึก)", level=1)
    doc.add_paragraph(f"ระบบ: {CFG.get('APP_TITLE', '-')}")
    doc.add_paragraph(f"วันที่สร้างเอกสาร: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # Resolve form values
    event_date_val = st.session_state.get("form_event_date", "")
    event_time_val = st.session_state.get("form_event_time", "")
    group_name = st.session_state.get("form_incident_group", "")
    unit_name = st.session_state.get("form_service_unit", "")
    sev = st.session_state.get("form_severity", "")

    if isinstance(event_date_val, date):
        event_date_text = event_date_val.isoformat()
    else:
        event_date_text = str(event_date_val)

    if isinstance(event_time_val, time):
        event_time_text = event_time_val.strftime("%H:%M")
    else:
        event_time_text = str(event_time_val)

    code_option = st.session_state.get("form_event_code_option", "")
    code, topic = parse_event_code_option(code_option)
    if code == "OTHER":
        code = str(st.session_state.get("form_event_code_other_code", "") or "").strip()
        topic = str(st.session_state.get("form_event_code_other_topic", "") or "").strip()

    doc.add_heading("1) ข้อมูลเหตุการณ์", level=2)
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"

    def add_row(k: str, v: str):
        row = t.add_row().cells
        row[0].text = str(k)
        row[1].text = str(v or "")

    add_row("หน่วย", unit_name)
    add_row("วันที่เกิดเหตุ", event_date_text)
    add_row("เวลาเกิดเหตุ", event_time_text)
    add_row("กลุ่มเหตุการณ์", group_name)
    add_row("รหัสเหตุการณ์", code)
    add_row("หัวข้อเหตุการณ์", topic)
    add_row("ระดับความรุนแรง", f"{sev} ({current_severity_scheme(group_name)})")
    sev_desc = severity_description(sev, group_name)
    if sev_desc:
        add_row("คำอธิบายระดับ", sev_desc)

    doc.add_heading("2) รายละเอียดเหตุการณ์", level=2)
    doc.add_paragraph(st.session_state.get("form_incident_detail", "") or "-")

    doc.add_heading("3) ข้อมูลเสริม (จากผู้ใช้)", level=2)
    doc.add_paragraph("3.1 ไทม์ไลน์")
    doc.add_paragraph(st.session_state.get("form_timeline_text", "") or "-")

    doc.add_paragraph("3.2 การแก้ไขเบื้องต้น")
    doc.add_paragraph(st.session_state.get("form_initial_correction", "") or "-")

    doc.add_paragraph("3.3 RCA (ข้อความ)")
    doc.add_paragraph(st.session_state.get("form_rca_text", "") or "-")

    doc.add_paragraph("3.4 แผนพัฒนา")
    doc.add_paragraph(st.session_state.get("form_development_plan", "") or "-")

    analysis = st.session_state.get("rca_analysis_json") or {}
    plan = st.session_state.get("rca_plan_json") or {}

    if analysis:
        doc.add_heading("4) ผลวิเคราะห์ RCA จากระบบ", level=2)

        doc.add_paragraph("4.1 สรุปเหตุการณ์")
        doc.add_paragraph(str(analysis.get("event_summary", "-")))

        timeline = analysis.get("timeline", []) or []
        doc.add_paragraph("4.2 ไทม์ไลน์เหตุการณ์")
        if timeline:
            for item in timeline:
                doc.add_paragraph(f"- {item}")
        else:
            doc.add_paragraph("-")

        fishbone = analysis.get("fishbone", {}) or {}
        doc.add_paragraph("4.3 Fishbone (สรุปแบบข้อความ)")
        effect = fishbone.get("effect", "")
        if effect:
            doc.add_paragraph(f"ผลลัพธ์/เหตุการณ์: {effect}")
        for cat in (fishbone.get("categories", []) or []):
            label = str(cat.get("label", "") or "ไม่ระบุ")
            doc.add_paragraph(f"หมวด: {label}")
            for it in (cat.get("items", []) or []):
                doc.add_paragraph(f"  - {it}")

        whys = analysis.get("five_whys", []) or []
        doc.add_paragraph("4.4 5 Whys")
        if whys:
            for w in whys:
                doc.add_paragraph(f"- {w}")
        else:
            doc.add_paragraph("-")

        swiss = analysis.get("swiss_cheese", []) or []
        doc.add_paragraph("4.5 Swiss Cheese")
        if swiss:
            for row in swiss:
                line = (
                    f"[{row.get('layer','')}] "
                    f"type={row.get('type','')} | "
                    f"hole={row.get('hole','')} | "
                    f"prevention={row.get('prevention','')}"
                )
                doc.add_paragraph(f"- {line}")
        else:
            doc.add_paragraph("-")

        factors = analysis.get("contributing_factors", []) or []
        doc.add_paragraph("4.6 ปัจจัยเอื้อ/ปัจจัยร่วม")
        if factors:
            for f in factors:
                doc.add_paragraph(f"- {f}")
        else:
            doc.add_paragraph("-")

    if plan:
        doc.add_heading("5) แผนปฏิบัติการ / PDSA จากระบบ", level=2)

        pdsa = plan.get("pdsa", {}) or {}
        for key_th, key_en in [("Plan", "plan"), ("Do", "do"), ("Study", "study"), ("Act", "act")]:
            doc.add_paragraph(f"PDSA - {key_th}")
            items = pdsa.get(key_en, []) or []
            if items:
                for it in items:
                    doc.add_paragraph(f"- {it}")
            else:
                doc.add_paragraph("-")

        ap = plan.get("action_plan", []) or []
        doc.add_paragraph("Action Plan")
        if ap:
            for i, row in enumerate(ap, 1):
                line = (
                    f"{i}) {row.get('measure','')} | "
                    f"ผู้รับผิดชอบ: {row.get('owner','')} | "
                    f"กำหนดเสร็จ: {row.get('due','')} | "
                    f"KPI: {row.get('kpi','')}"
                )
                doc.add_paragraph(line)
        else:
            doc.add_paragraph("-")

        ideas = plan.get("initiative_ideas", {}) or {}
        doc.add_paragraph("Initiative Ideas - Quick Wins (0–30 วัน)")
        for x in ideas.get("quick_wins_0_30_days", []) or []:
            doc.add_paragraph(f"- {x}")

        doc.add_paragraph("Initiative Ideas - ระยะกลาง (1–3 เดือน)")
        for x in ideas.get("mid_term_1_3_months", []) or []:
            doc.add_paragraph(f"- {x}")

        doc.add_paragraph("Initiative Ideas - ระยะยาว (3–12 เดือน)")
        for x in ideas.get("long_term_3_12_months", []) or []:
            doc.add_paragraph(f"- {x}")

        recs = plan.get("conclusion_recommendations", []) or []
        doc.add_paragraph("Conclusion & Recommendations")
        for i, x in enumerate(recs, 1):
            doc.add_paragraph(f"{i}. {x}")

        next72 = plan.get("next_72_hours", []) or []
        doc.add_paragraph("ก้าวถัดไป (ภายใน 72 ชั่วโมง)")
        for x in next72:
            doc.add_paragraph(f"- {x}")

    if uploaded_rca_image is not None:
        try:
            doc.add_heading("6) ภาพประกอบที่แนบ", level=2)
            img_bytes = uploaded_rca_image.getvalue()
            doc.add_paragraph(f"ชื่อไฟล์: {getattr(uploaded_rca_image, 'name', '-')}")
            doc.add_picture(BytesIO(img_bytes), width=Inches(6.2))
        except Exception as e:
            doc.add_paragraph(f"(ไม่สามารถแทรกรูปลง DOCX ได้: {e})")

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()


# =========================
# GEMINI / RCA ASSISTANT
# =========================

def call_gemini_json(
    prompt: str,
    api_key: str,
    image_file: Optional[Any] = None,
    timeout_sec: int = 60,
) -> Dict[str, Any]:
    if not api_key:
        raise ValueError("ยังไม่ได้ตั้งค่า GEMINI_API_KEY ใน Environment Variables")

    url = (
        "https://generativelanguage.googleapis.com/v1beta/models/"
        f"gemini-2.5-flash:generateContent?key={api_key}"
    )

    parts: List[Dict[str, Any]] = [{"text": prompt}]

    if image_file is not None:
        try:
            import base64
            img_bytes = image_file.getvalue()
            mime_type = getattr(image_file, "type", None) or "image/png"
            parts.append(
                {
                    "inline_data": {
                        "mime_type": mime_type,
                        "data": base64.b64encode(img_bytes).decode("utf-8"),
                    }
                }
            )
        except Exception:
            pass

    payload = {
        "contents": [{"parts": parts}],
        "generationConfig": {"responseMimeType": "application/json"},
        "safetySettings": [
            {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"},
        ],
    }

    resp = requests.post(url, json=payload, timeout=timeout_sec)
    try:
        data = resp.json()
    except Exception:
        raise RuntimeError(f"Gemini API ตอบกลับไม่ใช่ JSON (HTTP {resp.status_code})")

    if not resp.ok:
        err_msg = data.get("error", {}).get("message", f"Gemini API error ({resp.status_code})")
        raise RuntimeError(err_msg)

    text = (
        data.get("candidates", [{}])[0]
        .get("content", {})
        .get("parts", [{}])[0]
        .get("text", "")
    )

    if not text:
        raise RuntimeError("Gemini ไม่ส่งผลลัพธ์กลับมา")

    cleaned = re.sub(r"^```(?:json)?\s*", "", text.strip(), flags=re.I)
    cleaned = re.sub(r"\s*```$", "", cleaned.strip())

    try:
        return json.loads(cleaned)
    except json.JSONDecodeError as e:
        raise RuntimeError(f"Gemini ส่ง JSON ไม่ถูกต้อง: {e}\n\nRaw response:\n{cleaned[:2000]}")


def build_analysis_prompt(incident_text: str) -> str:
    return f"""
คุณคือผู้เชี่ยวชาญด้านความปลอดภัยผู้ป่วยและ RCA ในโรงพยาบาล
โปรดวิเคราะห์เหตุการณ์ต่อไปนี้เป็นภาษาไทย และส่งกลับเป็น JSON เท่านั้น (ห้ามมี markdown ห้ามมีข้อความอื่นนอก JSON)

เหตุการณ์:
\"\"\"{incident_text}\"\"\"

โครงสร้าง JSON ที่ต้องการ:
{{
  "event_summary": "สรุปเหตุการณ์แบบกระชับ 2-4 บรรทัด",
  "timeline": [
    "เหตุการณ์ลำดับที่ 1 ...",
    "เหตุการณ์ลำดับที่ 2 ..."
  ],
  "fishbone": {{
    "effect": "เหตุการณ์/ผลลัพธ์สั้นๆ",
    "categories": [
      {{
        "label": "คน",
        "items": ["...", "..."]
      }},
      {{
        "label": "วิธีการ",
        "items": ["...", "..."]
      }}
    ]
  }},
  "five_whys": [
    "ทำไม 1: ...",
    "ทำไม 2: ...",
    "ทำไม 3: ...",
    "ทำไม 4: ...",
    "ทำไม 5: ... (รากสาเหตุ)"
  ],
  "swiss_cheese": [
    {{
      "layer": "นโยบายองค์กร",
      "type": "latent/active",
      "hole": "ช่องโหว่",
      "prevention": "ข้อเสนอป้องกัน"
    }}
  ],
  "contributing_factors": [
    "ปัจจัยเอื้อ 1",
    "ปัจจัยเอื้อ 2"
  ]
}}

ข้อกำหนด:
- fishbone.categories มีได้สูงสุด 6 หมวด
- แต่ละหมวด items สูงสุด 5 ข้อ
- swiss_cheese อย่างน้อย 4 แถว
- five_whys ให้ครบ 5 ข้อ
- ใช้ภาษาไทยล้วน
    """.strip()


def build_plan_prompt(incident_text: str, analysis_json: Dict[str, Any]) -> str:
    analysis_text = json.dumps(analysis_json, ensure_ascii=False)
    return f"""
คุณคือผู้จัดการความปลอดภัยของโรงพยาบาล
จากเหตุการณ์และผลวิเคราะห์ RCA ด้านล่าง โปรดสร้างแผนปฏิบัติการ และส่งกลับเป็น JSON เท่านั้น

เหตุการณ์:
\"\"\"{incident_text}\"\"\"

ผลวิเคราะห์:
{analysis_text}

โครงสร้าง JSON:
{{
  "pdsa": {{
    "plan": ["...","..."],
    "do": ["...","..."],
    "study": ["...","..."],
    "act": ["...","..."]
  }},
  "action_plan": [
    {{
      "measure": "มาตรการ",
      "owner": "ผู้รับผิดชอบ",
      "due": "กำหนดเสร็จ",
      "kpi": "ตัวชี้วัด",
      "risk_control": "ความเสี่ยงและแนวทางลดเสี่ยง"
    }}
  ],
  "initiative_ideas": {{
    "quick_wins_0_30_days": ["...","..."],
    "mid_term_1_3_months": ["...","..."],
    "long_term_3_12_months": ["...","..."]
  }},
  "conclusion_recommendations": [
    "ข้อเสนอแนะสำคัญข้อ 1",
    "ข้อเสนอแนะสำคัญข้อ 2",
    "ข้อเสนอแนะสำคัญข้อ 3",
    "ข้อเสนอแนะสำคัญข้อ 4",
    "ข้อเสนอแนะสำคัญข้อ 5"
  ],
  "next_72_hours": [
    "ก้าวถัดไปภายใน 72 ชั่วโมง ข้อ 1",
    "ก้าวถัดไปภายใน 72 ชั่วโมง ข้อ 2"
  ]
}}

ข้อกำหนด:
- action_plan 3-8 แถว
- recommendation ให้ 5 ข้อพอดี
- ใช้ภาษาไทย
    """.strip()


# =========================
# RENDER ANALYSIS / PLAN
# =========================

def render_analysis_result(analysis: Dict[str, Any]):
    st.subheader("🔎 ผลวิเคราะห์ RCA")

    st.markdown("### 1) สรุปเหตุการณ์")
    st.write(analysis.get("event_summary", "-"))

    st.markdown("### 2) ไทม์ไลน์เหตุการณ์")
    timeline = analysis.get("timeline", []) or []
    if timeline:
        for i, item in enumerate(timeline, 1):
            st.markdown(f"- **{i}.** {item}")
    else:
        st.write("-")

    st.markdown("### 3) แผนผังก้างปลา (Ishikawa) — รายละเอียด")
    fishbone = analysis.get("fishbone", {}) or {}
    effect = fishbone.get("effect", "") or analysis.get("event_summary", "เหตุการณ์ / ผลลัพธ์")
    categories = fishbone.get("categories", []) or []

    st.markdown("**เหตุการณ์ / ผลลัพธ์**")
    st.write(effect if str(effect).strip() else "-")

    if categories:
        for idx, c in enumerate(categories, 1):
            label = str(c.get("label", "") or "ไม่ระบุ").strip()
            items = [str(x).strip() for x in (c.get("items", []) or []) if str(x).strip()]
            st.markdown(f"**{idx}) {label}**")
            if items:
                for item in items:
                    st.markdown(f"- {item}")
            else:
                st.markdown("- ไม่มีรายละเอียด")
    else:
        st.write("-")

    st.markdown("### 4) วิเคราะห์ทำไม-ทำไม (5 Whys)")
    whys = analysis.get("five_whys", []) or []
    if whys:
        for i, w in enumerate(whys, 1):
            st.markdown(f"{i}. {w}")
    else:
        st.write("-")

    st.markdown("### 5) โมเดลสวิสชีส")
    swiss = analysis.get("swiss_cheese", []) or []
    if swiss:
        df_swiss = pd.DataFrame(swiss).rename(
            columns={
                "layer": "ชั้นระบบ",
                "type": "ประเภท",
                "hole": "รู (ช่องโหว่)",
                "prevention": "มาตรการป้องกัน",
            }
        )
        st.dataframe(df_swiss, use_container_width=True, hide_index=True)
    else:
        st.write("-")

    factors = analysis.get("contributing_factors", []) or []
    if factors:
        st.markdown("### 6) ปัจจัยเอื้อ/ปัจจัยร่วม")
        for f in factors:
            st.markdown(f"- {f}")


def render_plan_result(plan: Dict[str, Any]):
    st.subheader("🎯 แผนปฏิบัติการ / PDSA")

    pdsa = plan.get("pdsa", {}) or {}
    pdsa_rows = [
        ["วางแผน (Plan)", "\n".join([f"- {x}" for x in (pdsa.get("plan", []) or [])])],
        ["ทำ (Do)", "\n".join([f"- {x}" for x in (pdsa.get("do", []) or [])])],
        ["ศึกษา (Study)", "\n".join([f"- {x}" for x in (pdsa.get("study", []) or [])])],
        ["ปรับปรุง (Act)", "\n".join([f"- {x}" for x in (pdsa.get("act", []) or [])])],
    ]
    st.markdown("### 1) PDSA")
    st.dataframe(
        pd.DataFrame(pdsa_rows, columns=["ขั้นตอน", "รายละเอียด"]),
        use_container_width=True,
        hide_index=True,
    )

    st.markdown("### 2) Action Plan")
    ap = plan.get("action_plan", []) or []
    if ap:
        df_ap = pd.DataFrame(ap).rename(
            columns={
                "measure": "มาตรการ",
                "owner": "ผู้รับผิดชอบ",
                "due": "กำหนดเสร็จ",
                "kpi": "KPI(ตัวชี้วัดผลลัพธ์)",
                "risk_control": "ความเสี่ยงและแนวทางลดเสี่ยง",
            }
        )
        st.dataframe(df_ap, use_container_width=True, hide_index=True)
    else:
        st.write("-")

    st.markdown("### 3) Initiative Ideas")
    ideas = plan.get("initiative_ideas", {}) or {}
    col1, col2, col3 = st.columns(3)
    with col1:
        st.markdown("**Quick Wins (0–30 วัน)**")
        for x in ideas.get("quick_wins_0_30_days", []) or []:
            st.markdown(f"- {x}")
    with col2:
        st.markdown("**ระยะกลาง (1–3 เดือน)**")
        for x in ideas.get("mid_term_1_3_months", []) or []:
            st.markdown(f"- {x}")
    with col3:
        st.markdown("**ระยะยาว (3–12 เดือน)**")
        for x in ideas.get("long_term_3_12_months", []) or []:
            st.markdown(f"- {x}")

    st.markdown("### 4) Conclusion & Recommendations")
    recs = plan.get("conclusion_recommendations", []) or []
    if recs:
        for i, x in enumerate(recs, 1):
            st.markdown(f"{i}. {x}")
    else:
        st.write("-")

    st.markdown("**ก้าวถัดไป (ภายใน 72 ชั่วโมง)**")
    next72 = plan.get("next_72_hours", []) or []
    if next72:
        for x in next72:
            st.markdown(f"- {x}")
    else:
        st.write("-")


# =========================
# FORM / SAVE
# =========================

def init_form_state_defaults():
    defaults = {
        "form_service_unit": UNIT_OPTIONS[0],
        "form_event_date": date.today(),
        "form_event_time": datetime.now().time().replace(second=0, microsecond=0),

        "form_incident_group": INCIDENT_GROUP_OPTIONS[0],
        "form_event_code_option": "",
        "form_event_code_other_code": "",
        "form_event_code_other_topic": "",

        "form_severity": "A",
        "form_drug_name": "",  # คงไว้เพื่อ compatibility (ไม่แสดงในฟอร์ม)
        "form_incident_detail": "",
        "form_timeline_text": "",
        "form_initial_correction": "",
        "form_rca_text": "",
        "form_development_plan": "",
        "rca_analysis_json": None,
        "rca_plan_json": None,
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

    # ensure event code option default matches group
    group_name = st.session_state.get("form_incident_group", INCIDENT_GROUP_OPTIONS[0])
    opts = event_code_options_for_group(group_name)
    if st.session_state.get("form_event_code_option", "") not in opts:
        st.session_state["form_event_code_option"] = opts[0] if opts else ""

    # ensure severity default matches scheme
    scheme = current_severity_scheme(group_name)
    sev = str(st.session_state.get("form_severity", ""))
    if scheme == "1-5" and sev not in SEVERITY_OPTIONS_PEOPLE:
        st.session_state["form_severity"] = "1"
    elif scheme == "A-I" and sev not in SEVERITY_OPTIONS_AI:
        st.session_state["form_severity"] = "A"


def validate_required_form() -> Tuple[bool, List[str]]:
    errs: List[str] = []

    if not str(st.session_state.get("form_service_unit", "")).strip():
        errs.append("กรุณาเลือกหน่วย")

    group_name = str(st.session_state.get("form_incident_group", "")).strip()
    if not group_name:
        errs.append("กรุณาเลือกกลุ่มเหตุการณ์")

    selected_option = str(st.session_state.get("form_event_code_option", "")).strip()
    code, _topic = parse_event_code_option(selected_option)
    if not selected_option:
        errs.append("กรุณาเลือกรหัสเหตุการณ์")
    elif code == "OTHER":
        other_code = str(st.session_state.get("form_event_code_other_code", "")).strip()
        other_topic = str(st.session_state.get("form_event_code_other_topic", "")).strip()
        if not other_code:
            errs.append("กรุณาระบุรหัสเหตุการณ์ (กรณีเลือก อื่น ๆ)")
        if not other_topic:
            errs.append("กรุณาระบุหัวข้อเหตุการณ์ (กรณีเลือก อื่น ๆ)")

    if not str(st.session_state.get("form_incident_detail", "")).strip():
        errs.append("กรุณากรอกรายละเอียดเหตุการณ์")

    sev = str(st.session_state.get("form_severity", "")).strip()
    if not sev:
        errs.append("กรุณาเลือกระดับความรุนแรง")

    return (len(errs) == 0, errs)


def create_record_from_form(
    uploaded_rca_image: Optional[Any],
    rca_image_drive_url: str = "",
) -> Dict[str, Any]:
    now = datetime.now()
    event_date_val = st.session_state.get("form_event_date")
    event_time_val = st.session_state.get("form_event_time")

    if isinstance(event_date_val, datetime):
        event_date_str = event_date_val.date().isoformat()
    elif isinstance(event_date_val, date):
        event_date_str = event_date_val.isoformat()
    else:
        event_date_str = str(event_date_val)

    if isinstance(event_time_val, datetime):
        event_time_str = event_time_val.strftime("%H:%M")
    elif isinstance(event_time_val, time):
        event_time_str = event_time_val.strftime("%H:%M")
    else:
        event_time_str = str(event_time_val)

    group_name = st.session_state.get("form_incident_group", "")
    selected_option = st.session_state.get("form_event_code_option", "")
    event_code, event_topic = parse_event_code_option(selected_option)

    if event_code == "OTHER":
        event_code = str(st.session_state.get("form_event_code_other_code", "")).strip()
        event_topic = str(st.session_state.get("form_event_code_other_topic", "")).strip()

    event_display = f"{event_code} | {event_topic}".strip(" |")
    sev_scheme = current_severity_scheme(group_name)

    record = {
        "record_id": now.strftime("%Y%m%d%H%M%S%f"),
        "unit_name": st.session_state.get("form_service_unit", "").strip(),  # เก็บ “หน่วย” ที่เลือก
        "app_title": CFG["APP_TITLE"],

        "event_date": event_date_str,
        "event_time": event_time_str,

        # compatibility fields
        "process_step": event_code,
        "drug_name": "",

        "severity_level": st.session_state.get("form_severity", ""),
        "incident_detail": st.session_state.get("form_incident_detail", "").strip(),
        "timeline_text": st.session_state.get("form_timeline_text", "").strip(),
        "initial_correction": st.session_state.get("form_initial_correction", "").strip(),
        "rca_text": st.session_state.get("form_rca_text", "").strip(),
        "rca_image_filename": getattr(uploaded_rca_image, "name", "") if uploaded_rca_image else "",
        "rca_image_drive_url": (rca_image_drive_url or "").strip(),
        "development_plan": st.session_state.get("form_development_plan", "").strip(),
        "created_at": now.isoformat(timespec="seconds"),
        "created_by": st.session_state.get("login_username", ""),

        # new fields
        "incident_group": group_name,
        "event_code": event_code,
        "event_topic": event_topic,
        "severity_scheme": sev_scheme,
        "event_display": event_display,
    }
    return record


def request_form_reset_after_save():
    st.session_state["_reset_form_after_save"] = True
    st.session_state["_save_success_message"] = "บันทึกข้อมูลสำเร็จ ✅"


def apply_pending_form_reset():
    if st.session_state.get("_reset_form_after_save", False):
        st.session_state["form_service_unit"] = UNIT_OPTIONS[0]
        st.session_state["form_event_date"] = date.today()
        st.session_state["form_event_time"] = datetime.now().time().replace(second=0, microsecond=0)
        st.session_state["form_incident_group"] = INCIDENT_GROUP_OPTIONS[0]

        first_opts = event_code_options_for_group(INCIDENT_GROUP_OPTIONS[0])
        st.session_state["form_event_code_option"] = first_opts[0] if first_opts else ""
        st.session_state["form_event_code_other_code"] = ""
        st.session_state["form_event_code_other_topic"] = ""

        st.session_state["form_severity"] = "A"
        st.session_state["form_drug_name"] = ""
        st.session_state["form_incident_detail"] = ""
        st.session_state["form_timeline_text"] = ""
        st.session_state["form_initial_correction"] = ""
        st.session_state["form_rca_text"] = ""
        st.session_state["form_development_plan"] = ""
        st.session_state["rca_analysis_json"] = None
        st.session_state["rca_plan_json"] = None
        st.session_state["show_fishbone_preview"] = False

        st.session_state.pop("form_rca_image", None)

        st.session_state["_reset_form_after_save"] = False


def render_event_selection_block():
    st.markdown("### ประเภทและรหัสเหตุการณ์")

    # ปุ่มเลือกกลุ่ม (ใช้ radio แบบ horizontal ให้คล้ายปุ่ม)
    prev_group = st.session_state.get("form_incident_group", INCIDENT_GROUP_OPTIONS[0])
    st.radio(
        "เลือกกลุ่มก่อนเข้าสู่กระบวนการ",
        options=INCIDENT_GROUP_OPTIONS,
        horizontal=True,
        key="form_incident_group",
    )
    curr_group = st.session_state.get("form_incident_group", INCIDENT_GROUP_OPTIONS[0])

    # หากเปลี่ยนกลุ่ม ให้ reset ตัวเลือกรหัส + severity ให้เหมาะ
    if curr_group != prev_group:
        opts = event_code_options_for_group(curr_group)
        st.session_state["form_event_code_option"] = opts[0] if opts else ""
        st.session_state["form_event_code_other_code"] = ""
        st.session_state["form_event_code_other_topic"] = ""

        if current_severity_scheme(curr_group) == "1-5":
            st.session_state["form_severity"] = "1"
        else:
            st.session_state["form_severity"] = "A"

    opts = event_code_options_for_group(curr_group)
    if st.session_state.get("form_event_code_option") not in opts:
        st.session_state["form_event_code_option"] = opts[0] if opts else ""

    st.selectbox(
        "ค้นหารหัสเหตุการณ์ / หัวข้อย่อย (พิมพ์เพื่อค้นหาได้)",
        options=opts,
        key="form_event_code_option",
        help="พิมพ์รหัส เช่น CPM201, CPP101, GOI102 เพื่อค้นหาได้เร็วขึ้น",
    )

    selected_opt = st.session_state.get("form_event_code_option", "")
    code, topic = parse_event_code_option(selected_opt)

    if code == "OTHER":
        c1, c2 = st.columns([1, 2.2])
        with c1:
            st.text_input("รหัสเหตุการณ์ (อื่น ๆ)", key="form_event_code_other_code", placeholder="เช่น XYZ999")
        with c2:
            st.text_input("หัวข้อเหตุการณ์ (อื่น ๆ)", key="form_event_code_other_topic", placeholder="ระบุหัวข้อเหตุการณ์เพิ่มเติม")
    else:
        st.markdown(
            f"<div class='helper-box'><div class='helper-title'>หัวข้อที่เลือก</div>"
            f"<div><b>{html.escape(code)}</b> — {html.escape(topic)}</div></div>",
            unsafe_allow_html=True,
        )


def render_entry_tab():
    init_form_state_defaults()
    apply_pending_form_reset()

    if st.session_state.get("_save_success_message"):
        st.success(st.session_state.pop("_save_success_message"))


    left, right = st.columns([1.18, 1], gap="large")

    uploaded_rca_image = None

    with left:
        st.markdown("### 📝 บันทึกข้อมูล")

        # ✅ เพิ่มช่อง “หน่วย” ไว้บนสุด
        st.selectbox("หน่วย", UNIT_OPTIONS, key="form_service_unit")

        # วันที่ / เวลา
        c1, c2 = st.columns(2)
        with c1:
            st.date_input("วันที่เกิดเหตุ", key="form_event_date")
        with c2:
            st.time_input("เวลาเกิดเหตุ", key="form_event_time")

        # กลุ่ม + รหัสเหตุการณ์
        render_event_selection_block()

        # ระดับความรุนแรง (dynamic)
        group_name = st.session_state.get("form_incident_group", INCIDENT_GROUP_OPTIONS[0])
        sev_options = severity_options_for_group(group_name)
        if st.session_state.get("form_severity") not in sev_options:
            st.session_state["form_severity"] = sev_options[0]

        st.selectbox(
            f"ระดับความรุนแรง ({current_severity_scheme(group_name)})",
            sev_options,
            key="form_severity",
        )

        sev_desc = severity_description(st.session_state.get("form_severity", ""), group_name)
        if sev_desc:
            st.info(f"**ระดับ {st.session_state.get('form_severity','')}**: {sev_desc}")

        render_severity_guide(group_name)

        # รายละเอียดเหตุการณ์
        st.text_area("รายละเอียดเหตุการณ์", height=140, key="form_incident_detail")

        # ✅ ย้าย uploader มาไว้ถัดจากรายละเอียดเหตุการณ์ + เปลี่ยนชื่อ label
        uploaded_rca_image = st.file_uploader(
            "แนบภาพประกอบ(หากมี)",
            type=["png", "jpg", "jpeg", "webp"],
            key="form_rca_image",
            help="หากแนบไฟล์และกดบันทึก ระบบจะเก็บชื่อไฟล์และลิงก์ Google Drive ในชีต",
        )

        if uploaded_rca_image is not None:
            st.image(
                uploaded_rca_image,
                caption=f"ภาพประกอบ: {uploaded_rca_image.name}",
                use_container_width=True,
            )

        st.markdown("---")
        st.markdown("### ข้อมูลเสริม (ก่อนบันทึก)")
        st.text_area("1) ไทม์ไลน์", height=120, key="form_timeline_text")
        st.text_area("2) การแก้ไขเบื้องต้น", height=100, key="form_initial_correction")
        st.text_area("3) RCA (ข้อความ)", height=180, key="form_rca_text")
        st.text_area("4) แผนพัฒนา", height=140, key="form_development_plan")

        st.markdown("---")

        try:
            docx_bytes = build_docx_report_bytes(uploaded_rca_image=uploaded_rca_image)
            st.download_button(
                "📄 ดาวน์โหลดรายงาน DOCX (ก่อนบันทึก)",
                data=docx_bytes,
                file_name=f"RCA_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        except Exception as e:
            st.caption(f"ยังไม่สามารถสร้าง DOCX ได้: {e}")

        if st.button("💾 บันทึกข้อมูล", type="primary", use_container_width=True):
            ok, errs = validate_required_form()
            if not ok:
                for e in errs:
                    st.error(e)
            else:
                try:
                    record = create_record_from_form(uploaded_rca_image=uploaded_rca_image)

                    if uploaded_rca_image is not None:
                        drive_meta = upload_rca_image_to_drive(
                            uploaded_rca_image,
                            record_id=record["record_id"],
                        )
                        record["rca_image_filename"] = drive_meta.get("file_name", "") or getattr(uploaded_rca_image, "name", "")
                        record["rca_image_drive_url"] = drive_meta.get("file_url", "") or ""

                    append_record_to_sheet(record)

                    try:
                        load_sheet_df.clear()
                    except Exception:
                        pass

                    request_form_reset_after_save()
                    st.rerun()

                except Exception as e:
                    st.exception(e)

    with right:
        st.markdown("### 🧸 RCA Assistant")
        st.caption("ระบบจะวิเคราะห์จากรายละเอียดเหตุการณ์ แล้วแสดงผลให้ตรวจทาน จากนั้นคัดลอกไปกรอกในฟอร์มเองก่อนบันทึก")

        st.info(
            "หลักการใช้งาน: ปุ่ม RCA Assistant จะ **ไม่บันทึกลง Google Sheets** โดยอัตโนมัติ\n"
            "→ ผู้ใช้ตรวจทานผลลัพธ์ ก่อนนำไปกรอกในฟอร์ม แล้วค่อยกด **บันทึกข้อมูล**"
        )

        if st.button("🧸 RCA Assistant", use_container_width=True):
            incident_text = st.session_state.get("form_incident_detail", "").strip()
            if not incident_text:
                st.warning("กรุณากรอกรายละเอียดเหตุการณ์ก่อน")
            else:
                try:
                    with st.spinner("กำลังวิเคราะห์ RCA..."):
                        analysis = call_gemini_json(
                            prompt=build_analysis_prompt(incident_text),
                            api_key=CFG["GEMINI_API_KEY"],
                            image_file=uploaded_rca_image,  # แนบภาพประกอบไปให้ AI ได้ (ถ้ามี)
                            timeout_sec=90,
                        )
                        plan = call_gemini_json(
                            prompt=build_plan_prompt(incident_text, analysis),
                            api_key=CFG["GEMINI_API_KEY"],
                            timeout_sec=90,
                        )

                        st.session_state.rca_analysis_json = analysis
                        st.session_state.rca_plan_json = plan

                    st.success("วิเคราะห์เสร็จแล้ว ✅")
                except Exception as e:
                    st.error(f"RCA Assistant error: {e}")

        analysis = st.session_state.get("rca_analysis_json")
        plan = st.session_state.get("rca_plan_json")

        if analysis:
            render_analysis_result(analysis)

        if plan:
            st.markdown("---")
            render_plan_result(plan)


# =========================
# HISTORY TAB
# =========================

def parse_event_datetime_columns(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()

    out["event_date"] = out.get("event_date", "").astype(str).str.strip()
    out["event_time"] = out.get("event_time", "").astype(str).str.strip()

    out["_event_date_dt"] = pd.to_datetime(out["event_date"], errors="coerce")
    out["_event_datetime"] = pd.to_datetime(
        out["event_date"].astype(str) + " " + out["event_time"].astype(str),
        errors="coerce",
    )
    out["_event_date_only"] = out["_event_date_dt"].dt.date
    return out


def render_history_tab():
    st.markdown("## 📚 ดูข้อมูลย้อนหลัง")

    try:
        df = load_sheet_df()
    except Exception as e:
        st.error(f"โหลดข้อมูลจาก Google Sheets ไม่สำเร็จ: {e}")
        return

    if df.empty:
        st.info("ยังไม่มีข้อมูลใน Google Sheets")
        return

    df = parse_event_datetime_columns(df)

    valid_dates_series = df["_event_date_dt"].dropna()
    if valid_dates_series.empty:
        min_d = date.today()
        max_d = date.today()
    else:
        min_d = valid_dates_series.min().date()
        max_d = valid_dates_series.max().date()

    if max_d < min_d:
        min_d, max_d = max_d, min_d

    st.markdown("### ตัวกรอง")
    c1, c2, c3, c4 = st.columns([1, 1, 1.1, 1.5])
    with c1:
        start_date = st.date_input("วันที่เริ่ม", value=min_d, key="hist_start")
    with c2:
        end_date = st.date_input("วันที่สิ้นสุด", value=max_d, key="hist_end")
    with c3:
        sev_selected = st.multiselect(
            "ระดับความรุนแรง",
            options=sorted([x for x in df["severity_level"].dropna().astype(str).unique() if str(x).strip()]),
            default=[],
            key="hist_sev",
        )
    with c4:
        keyword = st.text_input("ค้นหา (รหัส/หัวข้อ/รายละเอียด)", key="hist_kw").strip()

    c5, c6 = st.columns([1, 2])
    with c5:
        unit_selected = st.multiselect(
            "หน่วย",
            options=sorted([x for x in df["unit_name"].dropna().astype(str).unique() if str(x).strip()]),
            default=[],
            key="hist_unit",
        )
    with c6:
        group_selected = st.multiselect(
            "กลุ่มเหตุการณ์",
            options=sorted([x for x in df["incident_group"].dropna().astype(str).unique() if str(x).strip()]),
            default=[],
            key="hist_group",
        )

    if start_date > end_date:
        st.warning("วันที่เริ่มมากกว่าวันที่สิ้นสุด ระบบจะสลับให้โดยอัตโนมัติ")
        start_date, end_date = end_date, start_date

    m = pd.Series(True, index=df.index)
    m &= df["_event_date_only"].notna()
    m &= (df["_event_date_only"] >= start_date) & (df["_event_date_only"] <= end_date)

    if sev_selected:
        m &= df["severity_level"].astype(str).isin(sev_selected)

    if unit_selected:
        m &= df["unit_name"].astype(str).isin(unit_selected)

    if group_selected:
        m &= df["incident_group"].astype(str).isin(group_selected)

    if keyword:
        kw = keyword.lower()
        m &= (
            df["event_code"].astype(str).str.lower().str.contains(kw, na=False)
            | df["event_topic"].astype(str).str.lower().str.contains(kw, na=False)
            | df["event_display"].astype(str).str.lower().str.contains(kw, na=False)
            | df["incident_detail"].astype(str).str.lower().str.contains(kw, na=False)
            | df["rca_text"].astype(str).str.lower().str.contains(kw, na=False)
            | df["development_plan"].astype(str).str.lower().str.contains(kw, na=False)
        )

    filtered = df[m].copy()

    filtered["_created_at_dt"] = pd.to_datetime(filtered.get("created_at", ""), errors="coerce")
    filtered = filtered.sort_values(
        by=["_event_datetime", "_created_at_dt"],
        ascending=False,
        na_position="last",
    )

    st.markdown(f"**ผลลัพธ์ทั้งหมด:** {len(filtered):,} รายการ")

    if not filtered.empty:
        s1, s2, s3 = st.columns(3)
        with s1:
            st.metric("จำนวนรายการ", f"{len(filtered):,}")
        with s2:
            st.metric(
                "จำนวนหน่วย",
                f"{filtered['unit_name'].astype(str).replace('', pd.NA).dropna().nunique():,}",
            )
        with s3:
            st.metric(
                "จำนวนกลุ่มเหตุการณ์",
                f"{filtered['incident_group'].astype(str).replace('', pd.NA).dropna().nunique():,}",
            )

    display_cols = [
        "event_date",
        "event_time",
        "unit_name",
        "incident_group",
        "event_code",
        "event_topic",
        "severity_level",
        "severity_scheme",
        "incident_detail",
        "timeline_text",
        "initial_correction",
        "rca_text",
        "rca_image_filename",
        "rca_image_drive_url",
        "development_plan",
        "created_at",
        "created_by",
    ]
    for c in display_cols:
        if c not in filtered.columns:
            filtered[c] = ""

    st.dataframe(
        filtered[display_cols],
        use_container_width=True,
        hide_index=True,
        column_config={
            "event_date": "วันที่",
            "event_time": "เวลา",
            "unit_name": "หน่วย",
            "incident_group": "กลุ่มเหตุการณ์",
            "event_code": "รหัสเหตุการณ์",
            "event_topic": "หัวข้อเหตุการณ์",
            "severity_level": "ระดับ",
            "severity_scheme": "สเกล",
            "incident_detail": "รายละเอียดเหตุการณ์",
            "timeline_text": "ไทม์ไลน์",
            "initial_correction": "การแก้ไขเบื้องต้น",
            "rca_text": "RCA (ข้อความ)",
            "rca_image_filename": "ไฟล์ภาพประกอบ",
            "rca_image_drive_url": "ลิงก์ภาพ (Drive)",
            "development_plan": "แผนพัฒนา",
            "created_at": "เวลาบันทึก",
            "created_by": "ผู้บันทึก",
        },
    )

    csv_bytes = filtered[display_cols].to_csv(index=False).encode("utf-8-sig")
    st.download_button(
        "⬇️ ดาวน์โหลดผลลัพธ์ (CSV)",
        data=csv_bytes,
        file_name=f"incident_history_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
        mime="text/csv",
        use_container_width=False,
    )

    with st.expander("🔍 ดูรายละเอียดรายรายการ (เลือกจาก 20 รายการล่าสุด)"):
        preview = filtered.head(20).copy()
        if preview.empty:
            st.write("ไม่มีข้อมูล")
        else:
            labels = []
            for _, r in preview.iterrows():
                labels.append(
                    f"{r.get('event_date','')} {r.get('event_time','')} | "
                    f"{r.get('unit_name','-')} | "
                    f"{r.get('event_code','-')} | ระดับ {r.get('severity_level','-')}"
                )
            selected_idx = st.selectbox(
                "เลือกเหตุการณ์",
                options=list(range(len(labels))),
                format_func=lambda i: labels[i],
                key="hist_detail_picker",
            )
            row = preview.iloc[int(selected_idx)]

            st.markdown("### ข้อมูลเหตุการณ์")
            st.write(f"**หน่วย:** {row.get('unit_name','')}")
            st.write(f"**กลุ่มเหตุการณ์:** {row.get('incident_group','')}")
            st.write(f"**รหัสเหตุการณ์:** {row.get('event_code','')}")
            st.write(f"**หัวข้อเหตุการณ์:** {row.get('event_topic','')}")
            st.write(f"**ระดับความรุนแรง:** {row.get('severity_level','')} ({row.get('severity_scheme','')})")

            st.markdown("### รายละเอียดเหตุการณ์")
            st.write(row.get("incident_detail", ""))

            st.markdown("### ไทม์ไลน์")
            st.write(row.get("timeline_text", ""))

            st.markdown("### การแก้ไขเบื้องต้น")
            st.write(row.get("initial_correction", ""))

            st.markdown("### RCA")
            st.write(row.get("rca_text", ""))

            drive_url = str(row.get("rca_image_drive_url", "")).strip()
            if drive_url:
                st.markdown("### ลิงก์ภาพประกอบ (Google Drive)")
                st.markdown(f"[เปิดไฟล์ภาพประกอบบน Google Drive]({drive_url})")

            st.markdown("### แผนพัฒนา")
            st.write(row.get("development_plan", ""))

            if str(row.get("rca_image_filename", "")).strip():
                st.caption(f"แนบไฟล์ไว้ตอนบันทึก: {row.get('rca_image_filename')}")

# =========================
# MAIN
# =========================

def render_header():
    return


def check_required_env():
    missing = []
    for key in ["GSHEET_URL", "GCP_SERVICE_ACCOUNT_JSON"]:
        if not CFG.get(key):
            missing.append(key)

    if missing:
        st.error("ยังตั้งค่า Environment Variables ไม่ครบ: " + ", ".join(missing))
        st.stop()

    if not str(CFG.get("GDRIVE_FOLDER_ID", "") or "").strip():
        st.warning(
            "ยังไม่ได้ตั้งค่า GDRIVE_FOLDER_ID → หากแนบภาพประกอบแล้วกดบันทึก ระบบจะอัปโหลดภาพไป Google Drive ไม่ได้"
        )


def main():
    ensure_auth_state()

    if not st.session_state.authenticated:
        render_login()
        return

    check_required_env()
    render_header()

    # CSS เฉพาะปุ่ม logout ให้เล็กลงและดูเนียนขึ้น
    st.markdown(
        """
        <style>
        .logout-mini-wrap {
            margin-top: 2px;
            text-align: right;
        }
        .logout-mini-wrap div[data-testid="stButton"] > button {
            padding: 0.25rem 0.55rem !important;
            min-height: 32px !important;
            font-size: 0.92rem !important;
            border-radius: 8px !important;
            white-space: nowrap !important;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    # แถวบนสุดสำหรับ Logout (ไม่เอา tabs ไว้ใน columns เพื่อเลี่ยง nested columns error)
    _, logout_col = st.columns([14, 1.2], gap="small")
    with logout_col:
        st.markdown('<div class="logout-mini-wrap">', unsafe_allow_html=True)
        if st.button("🚪 Logout", key="logout_inline"):
            st.session_state.authenticated = False
            st.session_state.login_username = ""
            st.session_state.show_fishbone_preview = False
            st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)

    # Tabs ต้องอยู่ top-level (สำคัญมาก)
    tab1, tab2 = st.tabs(["บันทึกข้อมูล", "ดูข้อมูลย้อนหลัง"])

    with tab1:
        render_entry_tab()

    with tab2:
        render_history_tab()


if __name__ == "__main__":
    main()
